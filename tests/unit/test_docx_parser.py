"""Tests for DOCX parser support in SimpleTextParser."""

from __future__ import annotations


def test_docx_parser_exists():
    """Verify DOCX mime type is in supported_mimes."""
    from ragbot.infrastructure.ocr.simple_text_parser import SimpleTextParser

    parser = SimpleTextParser()
    mimes = parser.supported_mimes()
    assert "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in mimes


def test_docx_parser_method_exists():
    """Verify _parse_docx method is defined on SimpleTextParser."""
    from ragbot.infrastructure.ocr.simple_text_parser import SimpleTextParser

    parser = SimpleTextParser()
    assert hasattr(parser, "_parse_docx")
    assert callable(parser._parse_docx)


def test_docx_parser_round_trip():
    """Create a minimal DOCX in-memory and verify parsing extracts text."""
    from docx import Document as DocxDocument
    import io

    from ragbot.infrastructure.ocr.simple_text_parser import SimpleTextParser

    # Create a minimal DOCX in memory
    doc = DocxDocument()
    doc.add_heading("Test Heading", level=1)
    doc.add_paragraph("This is a test paragraph.")
    doc.add_heading("Sub Heading", level=2)
    doc.add_paragraph("Another paragraph here.")

    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    parser = SimpleTextParser()
    result = parser._parse_docx(docx_bytes)

    assert "# Test Heading" in result
    assert "## Sub Heading" in result
    assert "This is a test paragraph." in result
    assert "Another paragraph here." in result


def test_docx_parser_table_extraction():
    """Verify DOCX table is converted to markdown format."""
    from docx import Document as DocxDocument
    import io

    from ragbot.infrastructure.ocr.simple_text_parser import SimpleTextParser

    doc = DocxDocument()
    doc.add_paragraph("Before table.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header1"
    table.cell(0, 1).text = "Header2"
    table.cell(1, 0).text = "Val1"
    table.cell(1, 1).text = "Val2"

    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    parser = SimpleTextParser()
    result = parser._parse_docx(docx_bytes)

    assert "Header1" in result
    assert "Header2" in result
    assert "Val1" in result
    assert "Val2" in result
    assert "|" in result  # markdown table format
    assert "---" in result  # separator row
