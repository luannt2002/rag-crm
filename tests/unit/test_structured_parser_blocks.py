"""Structured parsers emit a typed Block stream via parse_blocks (A-I5 / B-2)."""

from __future__ import annotations

import io

import pytest

from ragbot.application.ports.document_parser_port import StructuredParserPort
from ragbot.domain.entities.document import Block


@pytest.mark.asyncio
async def test_docx_parse_blocks_returns_typed_blocks() -> None:
    from docx import Document as DocxDocument

    from ragbot.infrastructure.parser.docx_parser import DocxParser

    doc = DocxDocument()
    doc.add_heading("Pricing", level=1)
    doc.add_paragraph("Intro prose line.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Price"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "100"
    buf = io.BytesIO()
    doc.save(buf)

    parser = DocxParser()
    assert isinstance(parser, StructuredParserPort)
    blocks = await parser.parse_blocks(buf.getvalue(), file_name="doc.docx")

    assert len(blocks) >= 1
    assert all(isinstance(b, Block) for b in blocks)
    types = {b.type for b in blocks}
    assert "HEADING" in types
    assert "TABLE" in types
    # Atomic invariant — heading + table are never cut.
    for b in blocks:
        if b.type in {"HEADING", "TABLE"}:
            assert b.is_atomic is True


@pytest.mark.asyncio
async def test_excel_parse_blocks_returns_typed_blocks() -> None:
    from openpyxl import Workbook

    from ragbot.infrastructure.parser.excel_openpyxl_parser import (
        ExcelOpenpyxlParser,
    )

    wb = Workbook()
    ws = wb.active
    ws.append(["Name", "Price"])
    ws.append(["A", "100"])
    ws.append(["B", "200"])
    buf = io.BytesIO()
    wb.save(buf)

    parser = ExcelOpenpyxlParser()
    assert isinstance(parser, StructuredParserPort)
    blocks = await parser.parse_blocks(buf.getvalue(), file_name="prices.xlsx")

    assert len(blocks) >= 1
    assert all(isinstance(b, Block) for b in blocks)
    assert any(b.type == "TABLE" and b.is_atomic for b in blocks)


@pytest.mark.asyncio
async def test_google_sheets_parse_blocks_returns_typed_blocks() -> None:
    from ragbot.infrastructure.parser.google_sheets_parser import (
        GoogleSheetsParser,
    )

    csv_bytes = b"Name,Price\nA,100\nB,200\n"
    parser = GoogleSheetsParser()
    assert isinstance(parser, StructuredParserPort)
    blocks = await parser.parse_blocks(csv_bytes, file_name="sheet.csv")

    assert len(blocks) >= 1
    assert all(isinstance(b, Block) for b in blocks)
    assert any(b.type == "TABLE" for b in blocks)


@pytest.mark.asyncio
async def test_empty_input_yields_empty_block_list() -> None:
    from ragbot.infrastructure.parser.google_sheets_parser import (
        GoogleSheetsParser,
    )

    parser = GoogleSheetsParser()
    assert await parser.parse_blocks(b"", file_name="empty.csv") == []
