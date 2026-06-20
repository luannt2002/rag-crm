"""All common formats resolve, via the production registry, to a parser that
emits STRUCTURED output — not flat text. Regression guard for the 2026-06-19
flat-PDF fix and the "all formats → structured markdown" coverage claim.

Tables are the discriminator: flat extraction collapses them, structured
extraction keeps the row/column relationship (markdown ``| … |`` for DOCX/HTML,
``Column: value`` linearisation for spreadsheets).
"""

from __future__ import annotations

from io import BytesIO

from ragbot.infrastructure.parser.registry import detect_parser

_DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
_XLSX_MIME = (
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
)


def _docx_bytes() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("Chuong Mot", level=1)
    doc.add_paragraph("noi dung gioi thieu")
    doc.add_heading("Dieu Mot", level=2)
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Cot A"
    table.rows[0].cells[1].text = "Cot B"
    table.rows[1].cells[0].text = "gia tri 1"
    table.rows[1].cells[1].text = "gia tri 2"
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _xlsx_bytes() -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(["Dich vu", "Gia"])
    ws.append(["Tri mun", "700000"])
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


async def test_docx_emits_markdown_headings_and_table() -> None:
    parser = detect_parser(_DOCX_MIME, ".docx")
    assert parser is not None
    assert parser.get_provider_name() == "docx"
    md = "\n".join(c["content"] for c in await parser.parse(_docx_bytes(), file_name="t.docx"))
    assert "# Chuong Mot" in md  # h1 → #
    assert "## Dieu Mot" in md  # h2 → ##
    assert "| Cot A | Cot B |" in md  # table kept as markdown, not flattened


async def test_xlsx_row_carries_its_column_names() -> None:
    parser = detect_parser(_XLSX_MIME, ".xlsx")
    assert parser is not None
    assert parser.get_provider_name() == "excel_openpyxl"
    md = "\n".join(c["content"] for c in await parser.parse(_xlsx_bytes(), file_name="t.xlsx"))
    # Expert tabular RAG: each value travels with its column header.
    assert "Dich vu" in md and "Tri mun" in md and "700000" in md


async def test_html_emits_markdown_heading_and_table() -> None:
    parser = detect_parser("text/html", ".html")
    assert parser is not None
    assert parser.get_provider_name() == "kreuzberg_markdown"
    html = (
        b"<html><body><h1>Bang gia</h1><table>"
        b"<tr><th>Dich vu</th><th>Gia</th></tr>"
        b"<tr><td>Tri mun</td><td>700000</td></tr></table></body></html>"
    )
    md = "\n".join(c["content"] for c in await parser.parse(html, file_name="t.html"))
    assert "# Bang gia" in md
    assert "| Dich vu | Gia" in md  # table preserved
