"""DocxParser strategy — unit tests.

Builds an in-memory .docx via python-docx, asserts heading-aware chunking
and table extraction shape.
"""

from __future__ import annotations

import importlib.util
from io import BytesIO

import pytest

from ragbot.application.ports.document_parser_port import DocumentParserPort


_PYTHON_DOCX = importlib.util.find_spec("docx") is not None
_DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


@pytest.mark.skipif(not _PYTHON_DOCX, reason="python-docx not installed")
def test_docx_parser_supports_mime_and_ext() -> None:
    from ragbot.infrastructure.parser.docx_parser import DocxParser

    p = DocxParser()
    assert p.supports(_DOCX_MIME, ".docx") is True
    assert p.supports("", ".docx") is True
    assert p.supports(_DOCX_MIME, "") is True
    assert p.supports("application/pdf", ".pdf") is False
    assert p.get_provider_name() == "docx"
    assert isinstance(p, DocumentParserPort)


def _build_docx() -> bytes:
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_heading("Section A", level=1)
    doc.add_paragraph("Body of section A.")
    doc.add_heading("Sub of A", level=2)
    doc.add_paragraph("Detail line one.")
    doc.add_heading("Section B", level=1)
    doc.add_paragraph("Body of section B.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "ColX"
    table.cell(0, 1).text = "ColY"
    table.cell(1, 0).text = "valX"
    table.cell(1, 1).text = "valY"
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.mark.skipif(not _PYTHON_DOCX, reason="python-docx not installed")
@pytest.mark.asyncio
async def test_docx_parser_emits_structured_markdown_table_in_place() -> None:
    from ragbot.infrastructure.parser.docx_parser import DocxParser

    parser = DocxParser()
    chunks = await parser.parse(_build_docx(), file_name="sample.docx")

    # ONE structured-markdown document (AdapChunk L1), tables in document order.
    assert len(chunks) == 1
    md = chunks[0]["content"]
    assert chunks[0]["metadata"]["format"] == "markdown"
    assert chunks[0]["metadata"]["parser"] == "docx"
    assert chunks[0]["metadata"]["file_name"] == "sample.docx"

    # Headings + bodies preserved with their levels.
    assert "# Section A" in md and "## Sub of A" in md and "# Section B" in md
    assert "Body of section A." in md and "Detail line one." in md
    assert "Body of section B." in md

    # The table is rendered as markdown UNDER its preceding heading (Section B),
    # NOT appended at the document end bound to the wrong section (B3 fix).
    assert "| ColX | ColY |" in md
    assert "| --- | --- |" in md
    assert "| valX | valY |" in md
    assert md.index("# Section B") < md.index("| ColX | ColY |"), (
        "table must stay under its in-document section, not moved to the end"
    )


@pytest.mark.skipif(not _PYTHON_DOCX, reason="python-docx not installed")
@pytest.mark.asyncio
async def test_docx_parser_rejects_corrupt_bytes() -> None:
    from ragbot.infrastructure.parser.docx_parser import DocxParser

    parser = DocxParser()
    with pytest.raises(ValueError, match="Invalid or corrupt DOCX"):
        await parser.parse(b"not a docx", file_name="bad.docx")


@pytest.mark.skipif(not _PYTHON_DOCX, reason="python-docx not installed")
@pytest.mark.asyncio
async def test_docx_parser_rejects_oversized() -> None:
    from ragbot.infrastructure.parser.docx_parser import DocxParser
    from ragbot.shared.constants import DEFAULT_DOCX_MAX_BYTES

    parser = DocxParser()
    too_big = b"PK\x03\x04" + b"\x00" * (DEFAULT_DOCX_MAX_BYTES + 1)
    with pytest.raises(ValueError, match="too large"):
        await parser.parse(too_big, file_name="huge.docx")
