"""Simple OCR parser — fetches URL or accepts bytes, splits to structured blocks.

Preserves document structure: headings, tables, and paragraph boundaries.
This is a pragmatic baseline. Production should plug in
Docling / Mistral OCR via the OCRPort interface.
"""

from __future__ import annotations

import asyncio
import io
import re
from html.parser import HTMLParser

import httpx
import structlog

from ragbot.application.ports.ocr_port import OCRPort, ParsedDocument
from ragbot.domain.entities.document import Block
from ragbot.shared.constants import (
    DEFAULT_DOCX_MAX_BYTES,
    DEFAULT_LANGUAGE,
    DEFAULT_PARSER_HTTP_TIMEOUT_S,
)
from ragbot.shared.context_buffer import attach_context_buffer

logger = structlog.get_logger(__name__)

# ── Vietnamese section markers ──────────────────────────────────────────
_VN_SECTION_RE = re.compile(
    r"^(?:Phần|Chương|Mục|Điều|Khoản|Tiết)\s+\d+",
    re.IGNORECASE,
)
# ── Numbered section patterns (1., 1.1, 1.1.1, I., II., A.) ─────────
_NUMBERED_SECTION_RE = re.compile(
    r"^(?:"
    r"\d+(?:\.\d+)*\.?\s+"  # 1. or 1.1 or 1.1.1
    r"|[IVXLC]+\.\s+"       # I. II. III.
    r"|[A-Z]\.\s+"          # A. B. C.
    r"|Chapter\s+\d+"       # Chapter 1
    r"|CHAPTER\s+\d+"       # CHAPTER 1
    r")",
    re.IGNORECASE,
)
# ── Table row: line with 2+ pipe characters ──────────────────────────
_PIPE_TABLE_RE = re.compile(r"\|.*\|")
# ── Table separator row: |---|---|
_TABLE_SEP_RE = re.compile(r"^\|[\s\-:]+(?:\|[\s\-:]+)+\|?\s*$")
# ── Markdown heading ─────────────────────────────────────────────────
_MD_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")


class SimpleTextParser(OCRPort):
    """Parser supporting plain text, Markdown, HTML, PDF (via pypdfium2), and DOCX."""

    def __init__(
        self,
        *,
        http_timeout_s: float = DEFAULT_PARSER_HTTP_TIMEOUT_S,
        heading_detection: bool = True,
        table_detection: bool = True,
    ) -> None:
        self._client = httpx.AsyncClient(timeout=http_timeout_s)
        self._heading_detection = heading_detection
        self._table_detection = table_detection

    async def parse(
        self,
        source: str | bytes,
        *,
        mime_type_hint: str | None = None,
    ) -> ParsedDocument:
        # Resolve bytes
        if isinstance(source, str):
            r = await self._client.get(source)
            r.raise_for_status()
            data = r.content
        else:
            data = source

        mime = mime_type_hint or ""
        is_pdf = mime.endswith("pdf") or data[:4] == b"%PDF"
        is_html = "html" in mime
        is_markdown = "markdown" in mime
        is_docx = (
            mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or data[:4] == b"PK\x03\x04"
            and mime_type_hint
            and "wordprocessing" in mime_type_hint
        )

        text = ""
        page_count = 1

        if is_docx:
            text = await asyncio.get_event_loop().run_in_executor(
                None, self._parse_docx, data,
            )
            page_count = 1
        elif is_pdf:
            text = await asyncio.get_event_loop().run_in_executor(
                None, _extract_pdf_text, data,
            )
            page_count = max(1, text.count("\f") + 1)
        elif is_html:
            raw = data.decode("utf-8", errors="ignore")
            text = _parse_html_to_structured(raw)
        else:
            text = data.decode("utf-8", errors="ignore")

        # For markdown content, preserve structure as-is
        if is_markdown or (not is_pdf and not is_html and _looks_like_markdown(text)):
            is_markdown = True

        # Build structured blocks
        blocks = self._build_blocks(text, is_markdown=is_markdown)

        # AdapChunk Layer 2 — populate context_before/after on atomic
        # blocks (TABLE/HEADING etc.) so retrieval matches surrounding
        # prose. No-op when ``context_buffer_atomic_enabled`` is OFF.
        blocks = attach_context_buffer(blocks)

        # Domain-neutral: parser does NOT detect language. The caller
        # (DocumentService) overrides ``language`` from caller-provided value
        # or per-bot ``bots.language`` column. Default = deployment default
        # (``DEFAULT_LANGUAGE``); multi-industry tenants override per-bot.
        return ParsedDocument(
            blocks=blocks,
            language=DEFAULT_LANGUAGE,
            page_count=page_count,
        )

    def _build_blocks(self, text: str, *, is_markdown: bool) -> list[Block]:
        """Split text into typed blocks preserving headings and tables."""
        lines = text.split("\n")
        blocks: list[Block] = []
        current_lines: list[str] = []
        current_type: str = "TEXT"

        def _flush() -> None:
            nonlocal current_lines, current_type
            content = "\n".join(current_lines).strip()
            if content:
                blocks.append(Block(
                    type=current_type,
                    content=content,
                    is_atomic=current_type == "TABLE",
                ))
            current_lines = []
            current_type = "TEXT"

        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # ── Table detection ──
            if self._table_detection and _PIPE_TABLE_RE.match(stripped):
                _flush()
                table_lines = []
                while i < len(lines) and _PIPE_TABLE_RE.match(lines[i].strip()):
                    table_lines.append(lines[i].strip())
                    i += 1
                # Ensure separator row exists after header
                table_md = _normalize_pipe_table(table_lines)
                blocks.append(Block(
                    type="TABLE",
                    content=table_md,
                    is_atomic=True,
                ))
                continue

            # ── Tab/multi-space tabular data detection ──
            if self._table_detection and _is_tabular_line(stripped) and not is_markdown:
                # Look ahead for consecutive tabular lines
                tab_lines = []
                j = i
                while j < len(lines) and _is_tabular_line(lines[j].strip()):
                    tab_lines.append(lines[j].strip())
                    j += 1
                if len(tab_lines) >= 2:
                    _flush()
                    table_md = _table_lines_to_markdown(tab_lines)
                    blocks.append(Block(
                        type="TABLE",
                        content=table_md,
                        is_atomic=True,
                    ))
                    i = j
                    continue

            # ── Heading detection ──
            if self._heading_detection:
                heading_level = _detect_heading_level(stripped)
                if heading_level is not None:
                    _flush()
                    # Normalize to markdown heading format
                    heading_text = _strip_heading_markers(stripped)
                    prefix = "#" * heading_level
                    blocks.append(Block(
                        type="HEADING",
                        content=f"{prefix} {heading_text}",
                        is_atomic=True,
                    ))
                    i += 1
                    continue

            # ── Blank line = paragraph boundary ──
            if not stripped:
                _flush()
                i += 1
                continue

            # ── Normal text ──
            current_lines.append(line)
            i += 1

        _flush()

        # Fallback: if no blocks, return the full text as one block
        if not blocks:
            return [Block(type="TEXT", content=text.strip() or "", is_atomic=False)]

        return blocks

    _MAX_DOCX_BYTES = DEFAULT_DOCX_MAX_BYTES

    def _parse_docx(self, file_bytes: bytes) -> str:
        """Parse DOCX file preserving headings and table structure."""
        if len(file_bytes) > self._MAX_DOCX_BYTES:
            raise ValueError(
                f"DOCX file too large: {len(file_bytes)} bytes "
                f"(max {self._MAX_DOCX_BYTES})"
            )
        try:
            from docx import Document as DocxDocument

            try:
                doc = DocxDocument(io.BytesIO(file_bytes))
            except Exception as exc:  # noqa: BLE001 — python-docx raises opaque PackageNotFoundError/KeyError/zipfile.BadZipFile depending on corruption mode; we normalize to ValueError with type annotation preserved.
                raise ValueError(
                    f"Invalid or corrupt DOCX file: {type(exc).__name__}: {str(exc)[:100]}"
                ) from exc
            parts: list[str] = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                style = para.style.name if para.style else ""
                if style.startswith("Heading"):
                    level = 1
                    if style[-1].isdigit():
                        level = int(style[-1])
                    parts.append(f"{'#' * level} {text}")
                else:
                    parts.append(text)
            # Tables
            for table in doc.tables:
                rows: list[str] = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append("| " + " | ".join(cells) + " |")
                if rows:
                    header = rows[0]
                    separator = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
                    table_md = header + "\n" + separator + "\n" + "\n".join(rows[1:])
                    parts.append(table_md)
            return "\n\n".join(parts)
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX parsing. Install with: pip install python-docx",
            )

    def supported_mimes(self) -> frozenset[str]:
        return frozenset(
            {
                "application/pdf",
                "text/plain",
                "text/markdown",
                "text/html",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            },
        )

    async def close(self) -> None:
        await self._client.aclose()


# ── PDF extraction ──────────────────────────────────────────────────────


def _extract_pdf_text(data: bytes) -> str:
    try:
        import pypdfium2 as pdfium

        pdf = pdfium.PdfDocument(io.BytesIO(data))
        try:
            pages: list[str] = []
            for i in range(len(pdf)):
                page = pdf[i]
                tp = page.get_textpage()
                page_text = tp.get_text_range()
                pages.append(page_text)
                page.close()
            # Use form-feed between pages for page boundary preservation
            return "\f".join(pages)
        finally:
            pdf.close()
    except ImportError:
        logger.warning("pypdfium2 not installed; returning empty text")
        return ""


# ── HTML parsing ────────────────────────────────────────────────────────


def _parse_html_to_structured(html: str) -> str:
    """Convert HTML to structured text preserving headings and tables."""
    class _StructuredHTMLParser(HTMLParser):
        def __init__(self) -> None:
            super().__init__()
            self._output: list[str] = []
            self._current_tag: str | None = None
            self._in_table = False
            self._table_rows: list[list[str]] = []
            self._current_row: list[str] = []
            self._current_cell: list[str] = []
            self._heading_level = 0

        def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
            tag = tag.lower()
            self._current_tag = tag
            if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                self._heading_level = int(tag[1])
            elif tag == "table":
                self._in_table = True
                self._table_rows = []
            elif tag == "tr":
                self._current_row = []
            elif tag in ("td", "th"):
                self._current_cell = []
            elif tag == "br":
                self._output.append("\n")
            elif tag == "p":
                self._output.append("\n\n")

        def handle_endtag(self, tag: str) -> None:
            tag = tag.lower()
            if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                self._heading_level = 0
            elif tag in ("td", "th"):
                self._current_row.append(" ".join(self._current_cell).strip())
                self._current_cell = []
            elif tag == "tr":
                self._table_rows.append(self._current_row)
                self._current_row = []
            elif tag == "table":
                self._in_table = False
                if self._table_rows:
                    self._output.append("\n\n")
                    self._output.append(_rows_to_markdown_table(self._table_rows))
                    self._output.append("\n\n")
                self._table_rows = []

        def handle_data(self, data: str) -> None:
            text = data.strip()
            if not text:
                return
            if self._in_table:
                self._current_cell.append(text)
            elif self._heading_level > 0:
                prefix = "#" * self._heading_level
                self._output.append(f"\n\n{prefix} {text}\n\n")
            else:
                self._output.append(text + " ")

        def get_text(self) -> str:
            return "".join(self._output).strip()

    parser = _StructuredHTMLParser()
    parser.feed(html)
    return parser.get_text()


def _rows_to_markdown_table(rows: list[list[str]]) -> str:
    """Convert list of rows (each a list of cell strings) to markdown table."""
    if not rows:
        return ""
    # Normalize column count
    max_cols = max(len(r) for r in rows)
    normalized = [r + [""] * (max_cols - len(r)) for r in rows]

    lines: list[str] = []
    # Header row
    header = normalized[0]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join("---" for _ in header) + " |")
    # Data rows
    for row in normalized[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


# ── Heading detection ───────────────────────────────────────────────────


def _detect_heading_level(line: str) -> int | None:
    """Detect if a line is a heading. Returns level (1-6) or None.

    Heuristics:
    - Markdown headers: # Title -> level 1
    - ALL CAPS short lines (< 80 chars, > 2 chars): likely heading
    - Numbered sections: 1., 1.1, Chapter 1
    - Vietnamese section markers: Phần, Chương, Mục, Điều
    """
    if not line or len(line) > 200:
        return None

    # Markdown headers
    md_match = _MD_HEADING_RE.match(line)
    if md_match:
        return len(md_match.group(1))

    # Vietnamese section markers (Phần = level 1, Chương = level 2, Mục/Điều = level 3)
    if _VN_SECTION_RE.match(line):
        lower = line.lower()
        if lower.startswith("phần"):
            return 1
        if lower.startswith("chương") or lower.startswith("chapter"):
            return 2
        return 3

    # ALL CAPS lines (short, meaningful text — not just numbers/symbols)
    alpha_chars = [c for c in line if c.isalpha()]
    if (
        len(alpha_chars) >= 3
        and len(line) < 80
        and line == line.upper()
        and any(c.isalpha() for c in line)
    ):
        return 1

    # Numbered sections: "1. Title" or "1.1 Title" — only if short enough
    if len(line) < 120 and _NUMBERED_SECTION_RE.match(line):
        # Depth determines level: 1. → 2, 1.1 → 3, 1.1.1 → 4
        num_match = re.match(r"^(\d+(?:\.\d+)*)", line)
        if num_match:
            depth = num_match.group(1).count(".") + 1
            return min(depth + 1, 6)  # offset by 1 so "1." = h2
        # Roman numeral or letter sections
        return 2

    return None


def _strip_heading_markers(line: str) -> str:
    """Remove markdown heading markers (#) from line text."""
    md_match = _MD_HEADING_RE.match(line)
    if md_match:
        return md_match.group(2).strip()
    return line.strip()


# ── Table detection & conversion ────────────────────────────────────────


def _is_tabular_line(line: str) -> bool:
    """Check if a line looks like tabular data (tab-separated or multi-space aligned)."""
    if not line or len(line) < 5:
        return False
    # Tab-separated with 2+ columns
    if "\t" in line and line.count("\t") >= 1:
        return True
    # Multiple groups of 2+ spaces separating text
    parts = re.split(r"\s{2,}", line.strip())
    return len(parts) >= 3


def _table_lines_to_markdown(lines: list[str]) -> str:
    """Convert lines that look like a table into markdown table format.

    Handles tab-separated and multi-space aligned data.
    """
    if not lines:
        return ""

    rows: list[list[str]] = []
    for line in lines:
        if "\t" in line:
            cells = [c.strip() for c in line.split("\t") if c.strip() or rows]
        else:
            cells = [c.strip() for c in re.split(r"\s{2,}", line.strip())]
        if cells:
            rows.append(cells)

    if not rows:
        return "\n".join(lines)

    return _rows_to_markdown_table(rows)


def _normalize_pipe_table(lines: list[str]) -> str:
    """Normalize pipe-delimited table lines, ensuring separator row exists."""
    if not lines:
        return ""

    # Already well-formed markdown table
    if len(lines) >= 2 and _TABLE_SEP_RE.match(lines[1]):
        return "\n".join(lines)

    # Parse cells from each line
    rows: list[list[str]] = []
    for line in lines:
        # Split by pipe, strip outer empties
        cells = [c.strip() for c in line.split("|")]
        # Remove leading/trailing empty cells from pipe boundaries
        if cells and not cells[0]:
            cells = cells[1:]
        if cells and not cells[-1]:
            cells = cells[:-1]
        rows.append(cells)

    if not rows:
        return "\n".join(lines)

    return _rows_to_markdown_table(rows)


# ── Markdown detection ──────────────────────────────────────────────────


def _looks_like_markdown(text: str) -> bool:
    """Heuristic: does the text look like it's already Markdown?"""
    lines = text.split("\n")[:50]  # Check first 50 lines
    md_indicators = 0
    for line in lines:
        stripped = line.strip()
        if _MD_HEADING_RE.match(stripped):
            md_indicators += 1
        elif stripped.startswith("- ") or stripped.startswith("* "):
            md_indicators += 1
        elif stripped.startswith("```"):
            md_indicators += 1
        elif _PIPE_TABLE_RE.match(stripped):
            md_indicators += 1
    return md_indicators >= 2


__all__ = ["SimpleTextParser"]
